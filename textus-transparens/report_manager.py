import csv
import json
import re
from pathlib import Path
from sqlalchemy.orm import Session
from sqlalchemy import func

from code_manager import get_db_engine
from models import Code, CodeAssignment, Extract, Source, Case, Memo, Theme, ThemeCode, ThemeCluster, Cluster, ClusterCode, JudgementNote, AuditLog

def export_to_excel(data: list[list], xlsx_path: Path):
    import pandas as pd
    from openpyxl.utils import get_column_letter
    
    if not data:
        return
        
    df = pd.DataFrame(data[1:], columns=data[0])
    with pd.ExcelWriter(str(xlsx_path), engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Report')
        worksheet = writer.sheets['Report']
        
        for idx, col in enumerate(df.columns):
            series = df[col]
            max_len = 0
            if len(series) > 0:
                max_len = series.astype(str).map(len).max()
            max_len = max((max_len, len(str(series.name)))) + 2
            max_len = min(max_len, 100)
            col_letter = get_column_letter(idx + 1)
            worksheet.column_dimensions[col_letter].width = max_len

def _add_formatted_runs(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)

def export_to_docx(md_content: str, docx_path: Path):
    import docx
    doc = docx.Document()
    
    lines = md_content.split('\n')
    table_rows = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_table and table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_rows):
                    for j, cell in enumerate(row):
                        if j < len(table.columns):
                            clean_cell = cell.replace('**', '').replace('\\|', '|')
                            table.cell(i, j).text = clean_cell
                table_rows = []
                in_table = False
            continue
            
        if line.startswith('|') and line.endswith('|'):
            if '---' in line:
                continue
            cells = line.split('|')
            row = [cell.strip() for cell in cells[1:-1]]
            if row:
                table_rows.append(row)
            in_table = True
            continue
            
        if in_table:
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_rows):
                    for j, cell in enumerate(row):
                        if j < len(table.columns):
                            clean_cell = cell.replace('**', '').replace('\\|', '|')
                            table.cell(i, j).text = clean_cell
                table_rows = []
            in_table = False
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            _add_formatted_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
            
    if in_table and table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                if j < len(table.columns):
                    clean_cell = cell.replace('**', '').replace('\\|', '|')
                    table.cell(i, j).text = clean_cell
                    
    doc.save(str(docx_path))

def export_to_pdf(docx_path: Path, pdf_path: Path):
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Failed to export PDF via COM: {e}")

def generate_codebook_report(project_dir: Path, format: str = "default") -> list[str]:
    engine = get_db_engine(project_dir)
    reports_dir = project_dir / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    out_paths = []
    
    with Session(engine) as session:
        codes = session.query(Code).filter(Code.status == 'active').all()
        
        children_map = {}
        top_level = []
        for c in codes:
            if c.parent_code_id:
                children_map.setdefault(c.parent_code_id, []).append(c)
            else:
                top_level.append(c)
                
        def print_code(code, level=1):
            md = f"{'#' * level} {code.name} (ID: {code.code_id})\n\n"
            if code.definition:
                md += f"**Definition:** {code.definition}\n\n"
            if code.inclusion_rules:
                md += f"**Inclusion Rules:** {code.inclusion_rules}\n\n"
            if code.exclusion_rules:
                md += f"**Exclusion Rules:** {code.exclusion_rules}\n\n"
            if code.boundary_rules:
                md += f"**Boundary Rules:** {code.boundary_rules}\n\n"
            
            for child in children_map.get(code.code_id, []):
                md += print_code(child, level + 1)
            return md

        content = "# Codebook Report\n\n"
        for c in top_level:
            content += print_code(c, 2)
            
        csv_rows = [["Code ID", "Name", "Parent ID", "Definition", "Inclusion Rules", "Exclusion Rules", "Boundary Rules"]]
        for c in codes:
            csv_rows.append([
                c.code_id, c.name, c.parent_code_id or "", 
                c.definition or "", c.inclusion_rules or "", 
                c.exclusion_rules or "", c.boundary_rules or ""
            ])
            
        base_path = reports_dir / "codebook_report"
        
        if format in ("default", "md"):
            p = base_path.with_suffix(".md")
            with open(p, "w", encoding="utf-8") as f:
                f.write(content)
            out_paths.append(str(p))
            
        if format == "csv":
            p = base_path.with_suffix(".csv")
            with open(p, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(csv_rows)
            out_paths.append(str(p))
            
        if format == "xlsx":
            p = base_path.with_suffix(".xlsx")
            export_to_excel(csv_rows, p)
            out_paths.append(str(p))
            
        if format in ("docx", "pdf"):
            docx_p = base_path.with_suffix(".docx")
            export_to_docx(content, docx_p)
            if format == "docx":
                out_paths.append(str(docx_p))
            elif format == "pdf":
                pdf_p = base_path.with_suffix(".pdf")
                export_to_pdf(docx_p, pdf_p)
                out_paths.append(str(pdf_p))
                
        log_audit(session, "generate_report", "Report", "codebook_report")
        session.commit()
            
    return out_paths

def generate_extracts_report(project_dir: Path, format: str = "default") -> list[str]:
    engine = get_db_engine(project_dir)
    reports_dir = project_dir / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    out_paths = []
    
    with Session(engine) as session:
        assignments = session.query(CodeAssignment).filter(CodeAssignment.status == 'active').all()
        
        md_content = "# Extracts Report\n\n"
        md_content += "| Assignment ID | Code | Source | Anchor | Extract Text |\n"
        md_content += "|---|---|---|---|---|\n"
        
        csv_rows = [["Assignment ID", "Code ID", "Code Name", "Source ID", "Anchor", "Extract Text"]]
        
        for a in assignments:
            extract = a.extract
            source = extract.source if extract else None
            code = a.code
            
            code_name = code.name if code else "Unknown"
            code_id = code.code_id if code else ""
            source_id = source.source_id if source else "Unknown"
            anchor = extract.anchor if extract else "Unknown"
            text_span = extract.text_span if extract else ""
            
            safe_text = text_span.replace("\n", " ").replace("|", "\\|")
            md_content += f"| {a.assignment_id} | {code_name} | {source_id} | {anchor} | {safe_text} |\n"
            
            csv_rows.append([a.assignment_id, code_id, code_name, source_id, anchor, text_span])
            
        base_path = reports_dir / "extracts_report"
        
        if format in ("default", "md"):
            p = base_path.with_suffix(".md")
            with open(p, "w", encoding="utf-8") as f:
                f.write(md_content)
            out_paths.append(str(p))
            
        if format in ("default", "csv"):
            p = base_path.with_suffix(".csv")
            with open(p, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(csv_rows)
            out_paths.append(str(p))
            
        if format == "xlsx":
            p = base_path.with_suffix(".xlsx")
            export_to_excel(csv_rows, p)
            out_paths.append(str(p))
            
        if format in ("docx", "pdf"):
            docx_p = base_path.with_suffix(".docx")
            export_to_docx(md_content, docx_p)
            if format == "docx":
                out_paths.append(str(docx_p))
            elif format == "pdf":
                pdf_p = base_path.with_suffix(".pdf")
                export_to_pdf(docx_p, pdf_p)
                out_paths.append(str(pdf_p))
            
        log_audit(session, "generate_report", "Report", "extracts_report")
        session.commit()
            
    return out_paths

def generate_matrix_report(project_dir: Path, format: str = "default") -> list[str]:
    engine = get_db_engine(project_dir)
    reports_dir = project_dir / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    out_paths = []
    
    with Session(engine) as session:
        codes = session.query(Code).filter(Code.status == 'active').all()
        cases = session.query(Case).all()
        sources = session.query(Source).all()
        
        source_to_case = {}
        for s in sources:
            if s.attributes and "case_id" in s.attributes:
                source_to_case[s.source_id] = s.attributes["case_id"]
                
        memos = session.query(Memo).filter(Memo.case_id.isnot(None), Memo.source_id.isnot(None)).all()
        for m in memos:
            source_to_case[m.source_id] = m.case_id
            
        matrix = {code.code_id: {case.case_id: 0 for case in cases} for code in codes}
        for code in codes:
            matrix[code.code_id]["No Case"] = 0
            
        assignments = session.query(CodeAssignment).filter(CodeAssignment.status == 'active').all()
        for a in assignments:
            if not a.code_id or a.code_id not in matrix:
                continue
            
            extract = a.extract
            if not extract: continue
            
            source_id = extract.source_id
            case_id = source_to_case.get(source_id)
            
            if case_id in matrix[a.code_id]:
                matrix[a.code_id][case_id] += 1
            else:
                matrix[a.code_id]["No Case"] += 1
                
        csv_rows = []
        header = ["Code / Case"] + [c.name for c in cases] + ["No Case"]
        csv_rows.append(header)
        
        for code in codes:
            row = [code.name]
            for case in cases:
                row.append(matrix[code.code_id][case.case_id])
            row.append(matrix[code.code_id]["No Case"])
            csv_rows.append(row)
            
        md_content = "# Matrix Report\n\n"
        md_content += "| " + " | ".join(header) + " |\n"
        md_content += "|" + "|".join(["---"] * len(header)) + "|\n"
        for row in csv_rows[1:]:
            md_content += "| " + " | ".join(str(x) for x in row) + " |\n"
            
        base_path = reports_dir / "matrix_report"
        
        if format == "md":
            p = base_path.with_suffix(".md")
            with open(p, "w", encoding="utf-8") as f:
                f.write(md_content)
            out_paths.append(str(p))
            
        if format in ("default", "csv"):
            p = base_path.with_suffix(".csv")
            with open(p, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(csv_rows)
            out_paths.append(str(p))
            
        if format == "xlsx":
            p = base_path.with_suffix(".xlsx")
            export_to_excel(csv_rows, p)
            out_paths.append(str(p))
            
        if format in ("docx", "pdf"):
            docx_p = base_path.with_suffix(".docx")
            export_to_docx(md_content, docx_p)
            if format == "docx":
                out_paths.append(str(docx_p))
            elif format == "pdf":
                pdf_p = base_path.with_suffix(".pdf")
                export_to_pdf(docx_p, pdf_p)
                out_paths.append(str(pdf_p))
            
        log_audit(session, "generate_report", "Report", "matrix_report")
        session.commit()
        
    return out_paths

def generate_theme_pack(project_dir: Path, theme_id: int, format: str = "default") -> list[str]:
    engine = get_db_engine(project_dir)
    reports_dir = project_dir / "reports" / "theme_packs"
    
    out_paths = []
    
    with Session(engine) as session:
        theme = session.query(Theme).filter_by(theme_id=theme_id).first()
        if not theme:
            raise ValueError(f"Theme {theme_id} not found.")
            
        safe_theme_name = "".join(c if c.isalnum() else "_" for c in theme.name)
        theme_dir = reports_dir / safe_theme_name
        theme_dir.mkdir(parents=True, exist_ok=True)
        
        md = f"# Theme Evidence Pack: {theme.name}\n\n"
        if theme.description:
            md += f"**Description:** {theme.description}\n\n"
            
        theme_codes = session.query(ThemeCode).filter_by(theme_id=theme_id).all()
        code_ids = [tc.code_id for tc in theme_codes]
        
        theme_clusters = session.query(ThemeCluster).filter_by(theme_id=theme_id).all()
        cluster_ids = [tc.cluster_id for tc in theme_clusters]
        clusters = session.query(Cluster).filter(Cluster.cluster_id.in_(cluster_ids)).all() if cluster_ids else []
        
        if cluster_ids:
            cluster_codes = session.query(ClusterCode).filter(ClusterCode.cluster_id.in_(cluster_ids)).all()
            for cc in cluster_codes:
                if cc.code_id not in code_ids:
                    code_ids.append(cc.code_id)
                    
        codes = session.query(Code).filter(Code.code_id.in_(code_ids)).all() if code_ids else []
        
        md += "## Linked Clusters\n\n"
        if clusters:
            for c in clusters:
                md += f"- **{c.name}**: {c.description}\n"
        else:
            md += "None.\n"
        md += "\n"
        
        md += "## Linked Codes\n\n"
        if codes:
            for c in codes:
                md += f"- **{c.name}** (ID: {c.code_id}): {c.definition}\n"
        else:
            md += "None.\n"
        md += "\n"
        
        memos = session.query(Memo).filter(Memo.theme_id == theme_id).all()
        md += "## Theme Memos\n\n"
        if memos:
            for m in memos:
                md += f"### Memo {m.memo_id}\n{m.text}\n\n"
        else:
            md += "None.\n\n"
            
        md += "## Evidence Extracts\n\n"
        csv_rows = [["Extract Source", "Extract Anchor", "Code", "Text Span", "Judgement Note", "Confidence"]]
        
        if code_ids:
            assignments = session.query(CodeAssignment).filter(
                CodeAssignment.code_id.in_(code_ids),
                CodeAssignment.status == 'active'
            ).all()
            
            if assignments:
                for a in assignments:
                    extract = a.extract
                    code = a.code
                    source = extract.source if extract else None
                    if not extract: continue
                    
                    source_id_str = str(source.source_id) if source else 'Unknown'
                    md += f"### Assignment {a.assignment_id} (Code: {code.name})\n"
                    md += f"- **Source:** {source_id_str}\n"
                    md += f"- **Anchor:** {extract.anchor}\n"
                    md += f"> {extract.text_span}\n\n"
                    
                    notes = session.query(JudgementNote).filter_by(extract_id=extract.extract_id, proposed_code_id=code.code_id).all()
                    note_strs = []
                    conf_strs = []
                    for n in notes:
                        md += f"**Judgement Note:** {n.rationale} (Confidence: {n.confidence})\n\n"
                        note_strs.append(n.rationale or "")
                        conf_strs.append(n.confidence or "")
                        
                    csv_rows.append([
                        source_id_str, extract.anchor, code.name, extract.text_span,
                        " | ".join(note_strs), " | ".join(conf_strs)
                    ])
            else:
                md += "No extracts found for the linked codes.\n\n"
        else:
            md += "No codes linked to fetch extracts.\n\n"
            
        base_path = theme_dir / "evidence"
        
        if format in ("default", "md"):
            p = base_path.with_suffix(".md")
            with open(p, "w", encoding="utf-8") as f:
                f.write(md)
            out_paths.append(str(p))
            
        if format == "csv":
            p = base_path.with_suffix(".csv")
            with open(p, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerows(csv_rows)
            out_paths.append(str(p))
            
        if format == "xlsx":
            p = base_path.with_suffix(".xlsx")
            export_to_excel(csv_rows, p)
            out_paths.append(str(p))
            
        if format in ("docx", "pdf"):
            docx_p = base_path.with_suffix(".docx")
            export_to_docx(md, docx_p)
            if format == "docx":
                out_paths.append(str(docx_p))
            elif format == "pdf":
                pdf_p = base_path.with_suffix(".pdf")
                export_to_pdf(docx_p, pdf_p)
                out_paths.append(str(pdf_p))
            
        log_audit(session, "generate_theme_pack", "Report", str(theme_id))
        session.commit()
        
    return out_paths

def log_audit(session: Session, action: str, entity_type: str, entity_id: str, details: dict = None, user_id: str = "default_user"):
    audit = AuditLog(
        action=action,
        entity_type=entity_type,
        entity_id=str(entity_id),
        user_id=user_id,
        details=details or {}
    )
    session.add(audit)
